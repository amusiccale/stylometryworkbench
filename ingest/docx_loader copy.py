"""
ingest/docx_loader.py

DOCX file ingestion.
"""

from pathlib import Path
from docx import Document


def load_docx(path: str | Path) -> str:
    """
    Load a DOCX file and return its text.

    Args:
        path: Path to the DOCX file.

    Returns:
        Plain-text contents of the document.
    """

    path = Path(path)

    document = Document(path)

##    paragraphs = [
##        paragraph.text
##        for paragraph in document.paragraphs
##    ]

    paragraphs = [
            p.text
            for p in document.paragraphs
            if p.text.strip()
        ]

    return "\n".join(paragraphs)


